"""
Clinical report generator using learned rules (v3 — schema-driven resolution).

Loads report_rules.json, processes extraction data through each section rule,
calls Together AI / Fireworks for narrative generation with hallucination guard,
and assembles a DOCX.

Field resolution uses a dynamic alias map built from the form schema at startup.
Confidence scores propagate through the pipeline: low-confidence fields are
annotated or excluded depending on thresholds.
"""

import json
import logging
import os
import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, Cm, RGBColor
from openai import OpenAI

logger = logging.getLogger(__name__)

RULES_PATH = Path(__file__).parent.parent.parent / "report_learning" / "outputs" / "rules" / "report_rules.json"
_DEFAULT_SCHEMA_PATH = Path(__file__).parent.parent.parent / "templates" / "orofacial_exam_schema.json"


# =============================================================================
# Confidence thresholds (Part 5)
# =============================================================================

CONFIDENCE_USE_NORMALLY = 0.85
CONFIDENCE_INCLUDE_NO_FLAG = 0.60
CONFIDENCE_UNCERTAIN = 0.40

_LEGALLY_SIGNIFICANT_FIELDS = frozenset({
    "date_of_injury", "case_number", "claim_number", "patient_dob",
    "p1_birth_date", "p5_date_of_injury", "p1_date",
})


# =============================================================================
# Static legal / cover-page / preamble text constants
# =============================================================================

_LEGAL_DEMAND_BOX = (
    "Demand is hereby made for service of all medical reports relating to this claim, "
    "pursuant to CCR 10635(c).\n\n"
    "Please note:\nIf payment of this bill is denied, we will pursue provisions under L.C. 4603.2\n\n"
    "Please note:\nLabor Code 5402 (b)(c), requires the employer to authorize all appropriate "
    "medical care up to $10,000 until the liability for the claimed injury is accepted or rejected. "
    "As of 6/01/04, Labor Code 5814 mandates a 25% penalty on the amount of payment unreasonably "
    "delayed (10% if self-imposed).\n"
    "Accordingly, it would be requested that the defendant please provide immediate payment.\n\n"
    "Please note:\nAny report necessary for the billing company to fulfill its business obligation "
    "to the insurance company should be secured from the insurance company."
)

_PREAMBLE_DISCLAIMER = (
    "THIS IS AN EXAMINATION REPORT. THIS REPORT WILL BE INCORPORATED INTO THE PRIMARY TREATING "
    "PHYSICIAN'S PERMANENT AND STATIONARY REPORT AND, AS SUCH, WILL BE RELIED UPON AND WILL BE "
    "PART OF THE PRIMARY TREATING PHYSICIAN'S MEDICAL-LEGAL OPINION. USE OF THE CONTENTS OF THIS "
    "REPORT AND ITS OPINIONS AND CONCLUSIONS ARE FOR PROVING OR DISPROVING A CONTESTED CLAIM."
)

_PREAMBLE_EXAM_INTRO = (
    "This examination is an extensive oral cranial examination in that one hour was spent in "
    "evaluating {patient_title} {patient_last_name}, one hour was spent in interpretation of "
    "the diagnostic findings, and issues of medical causation are discussed."
)

_PREAMBLE_CONCLUSIONS_BASIS = (
    "I base my conclusions on the history given by the patient, clinical examination of the "
    "patient, diagnostic tests, and review of any medical records made available at the time "
    "of examination."
)

_PREAMBLE_RECORDS_DEMAND = (
    "Demand is hereby made for service of all medical reports and records relating to this "
    "claim, pursuant to CCR 10635(c), to be sent to my office for review and comment, as "
    "necessary. If no medical records or reports are forwarded for my review, then the parties "
    "cannot claim that my reporting is not considered substantial medical evidence by virtue of "
    "disregarding this demand."
)

_PREAMBLE_RESERVE_RIGHTS = (
    "I hereby reserve the right to change and/or alter any of my opinions, statements, and "
    "conclusions stated herein, following the provision of any new medical and/or dental "
    "records not provided at the time of examination."
)

_COVER_SKIP_SECTION_IDS = frozenset({
    "initial_report_in_the_field_of",
    "and_request_for_authorization",
})


# =============================================================================
# Schema-driven field alias map + legacy aliases
# =============================================================================

_LEGACY_ALIASES: dict[str, str] = {
    "Diagnostic_Bite_Force_Analysis_Left": "p16_left_newtons",
    "Diagnostic_Bite_Force_Analysis_Right": "p16_right_newtons",
    "p17_m26_69_osteoarthritis": "p17_m26_69_osteoarthrosis",
    "p7_tx_received": "p8_tx_received",
}


def _build_field_map_from_schema(schema_path: Path | None = None) -> dict[str, str]:
    """Build field aliases dynamically from the form schema."""
    path = schema_path or _DEFAULT_SCHEMA_PATH
    if not path.exists():
        logger.warning("Schema not found at %s — field alias map will be empty", path)
        return {}

    schema = json.loads(path.read_text(encoding="utf-8"))
    field_map: dict[str, str] = {}

    for page in schema.get("pages", []):
        for field in page.get("standalone_fields", []):
            _register_field(field, field_map)
        for section in page.get("sections", []):
            for field in section.get("fields", []):
                _register_field(field, field_map)
            for sub in section.get("subsections", []):
                for field in sub.get("fields", []):
                    _register_field(field, field_map)
        for table in page.get("tables", []):
            for col in table.get("columns", []):
                if col.get("column_id"):
                    col_id = col["column_id"]
                    header = col.get("header", "")
                    if header:
                        label_key = header.lower().replace(" ", "_")
                        field_map.setdefault(label_key, f"{table['table_id']}_{col_id}")

    return field_map


def _register_field(field: dict, field_map: dict[str, str]) -> None:
    fid = field.get("field_id", "")
    label = field.get("field_label", "")
    if not fid:
        return
    if label:
        label_key = label.lower().replace(" ", "_").replace("(", "").replace(")", "")
        field_map.setdefault(label_key, fid)
    if "_" in fid:
        short = fid.split("_", 1)[1]
        field_map.setdefault(short, fid)


_FIELD_MAP: dict[str, str] = _build_field_map_from_schema()


# =============================================================================
# Field value helpers
# =============================================================================

def _unwrap_field_value(fv: Any) -> Any:
    """Extract a scalar from a nested extraction field dict."""
    if not isinstance(fv, dict):
        return fv
    val = fv.get("value")
    if val is not None and val != "":
        return val
    checked = fv.get("is_checked")
    if checked is not None:
        return "Yes" if checked else "No"
    options = fv.get("circled_options")
    if options:
        return ", ".join(str(o) for o in options) if isinstance(options, list) else str(options)
    return val


def _get_confidence(fv: Any) -> float:
    """Extract the confidence score from a field value (1.0 for non-dict values)."""
    if isinstance(fv, dict):
        return float(fv.get("confidence", 1.0))
    return 1.0


def _flatten_all_fields(pages: list[dict]) -> dict[str, Any]:
    """Flatten all pages' field_values into a single dict."""
    all_fields: dict[str, Any] = {}
    for page in pages:
        for fid, fv in (page.get("field_values") or {}).items():
            if fid.startswith("__"):
                continue
            all_fields[fid] = fv
    return all_fields


def _resolve_field(
    field_id: str,
    all_fields: dict[str, Any],
    derived: dict[str, str],
    field_map: dict[str, str] | None = None,
) -> Any:
    """Resolve a field value through a 4-level lookup chain."""
    fm = field_map if field_map is not None else _FIELD_MAP

    fv = all_fields.get(field_id)
    if fv is not None:
        return _unwrap_field_value(fv)

    val = derived.get(field_id)
    if val is not None:
        return val

    fid_lower = field_id.lower().replace(" ", "_")
    canonical = fm.get(field_id) or fm.get(fid_lower) or _LEGACY_ALIASES.get(field_id)
    if canonical and canonical != field_id:
        fv = all_fields.get(canonical)
        if fv is not None:
            return _unwrap_field_value(fv)

    for key in all_fields:
        if key.lower() == fid_lower:
            return _unwrap_field_value(all_fields[key])

    return None


def _resolve_field_with_confidence(
    field_id: str,
    all_fields: dict[str, Any],
    field_map: dict[str, str] | None = None,
) -> tuple[Any, float]:
    """Resolve a field and return (unwrapped_value, confidence)."""
    fm = field_map if field_map is not None else _FIELD_MAP

    fv = all_fields.get(field_id)
    if fv is not None:
        return _unwrap_field_value(fv), _get_confidence(fv)

    fid_lower = field_id.lower().replace(" ", "_")
    canonical = fm.get(field_id) or fm.get(fid_lower) or _LEGACY_ALIASES.get(field_id)
    if canonical and canonical != field_id:
        fv = all_fields.get(canonical)
        if fv is not None:
            return _unwrap_field_value(fv), _get_confidence(fv)

    for key in all_fields:
        if key.lower() == fid_lower:
            fv = all_fields[key]
            return _unwrap_field_value(fv), _get_confidence(fv)

    return None, 1.0


def _collect_low_confidence_fields(all_fields: dict[str, Any]) -> list[dict]:
    """Scan all_fields for values below CONFIDENCE_INCLUDE_NO_FLAG."""
    results = []
    for fid, fv in all_fields.items():
        if fid.startswith("__"):
            continue
        conf = _get_confidence(fv)
        if conf < CONFIDENCE_INCLUDE_NO_FLAG:
            val = _unwrap_field_value(fv)
            if val is not None:
                results.append({"field_id": fid, "value": str(val), "confidence": round(conf, 3)})
    return results


# =============================================================================
# LLM client setup
# =============================================================================

TOGETHER_MODEL = "meta-llama/Llama-4-Maverick-17B-128E-Instruct-FP8"
TOGETHER_BASE_URL = "https://api.together.xyz/v1"
FIREWORKS_MODEL = "accounts/fireworks/models/deepseek-v3p1"
FIREWORKS_BASE_URL = "https://api.fireworks.ai/inference/v1"


def _build_llm_clients() -> list[tuple[OpenAI, str]]:
    """Build LLM clients in priority order for automatic fallback."""
    clients = []
    together_key = os.getenv("TOGETHER_API_KEY")
    fireworks_key = os.getenv("FIREWORKS_API_KEY")
    if together_key:
        clients.append((OpenAI(api_key=together_key, base_url=TOGETHER_BASE_URL, timeout=120.0), TOGETHER_MODEL))
    if fireworks_key:
        clients.append((OpenAI(api_key=fireworks_key, base_url=FIREWORKS_BASE_URL, timeout=120.0), FIREWORKS_MODEL))
    return clients


def _load_rules(rules_path: Path | None = None) -> dict[str, Any]:
    path = rules_path or RULES_PATH
    return json.loads(path.read_text(encoding="utf-8"))


# =============================================================================
# Conditions, templates, lists
# =============================================================================

def _check_conditions(
    conditions: list[dict],
    all_fields: dict[str, Any],
    derived: dict[str, str],
    field_map: dict[str, str],
) -> bool:
    """Evaluate section conditions against extraction data."""
    if not conditions:
        return True

    results = []
    for cond in conditions:
        fid = cond.get("field_id", "")
        op = cond.get("operator", "exists")
        expected = cond.get("value")
        val = _resolve_field(fid, all_fields, derived, field_map)

        if op == "exists":
            results.append(val is not None)
        elif op == "not_empty":
            results.append(val is not None and str(val).strip() != "")
        elif op == "equals":
            results.append(str(val).lower().strip() == str(expected).lower().strip() if val else False)
        elif op == "not_equals":
            results.append(str(val).lower().strip() != str(expected).lower().strip() if val else True)
        elif op == "contains":
            results.append(str(expected).lower() in str(val).lower() if val else False)
        elif op == "greater_than":
            try:
                results.append(float(val) > float(expected) if val else False)
            except (ValueError, TypeError):
                results.append(False)
        else:
            results.append(True)

    combine = conditions[0].get("combine", "and") if conditions else "and"
    return any(results) if combine == "or" else all(results)


def _fill_template(
    template: str,
    all_fields: dict[str, Any],
    derived: dict[str, str],
    field_map: dict[str, str],
) -> str:
    """Replace {field_id} and {field_id:default} placeholders.

    Legally significant fields with confidence < 0.60 are wrapped with
    [VERIFY: ...] so clinicians notice them immediately.
    """
    def replacer(match):
        token = match.group(1)
        parts = token.split(":")
        fid = parts[0]
        default = parts[1] if len(parts) > 1 else ""
        val, conf = _resolve_field_with_confidence(fid, all_fields, field_map)
        if val is None:
            val = derived.get(fid)
            conf = 1.0
        if val is None:
            return default
        if fid in _LEGALLY_SIGNIFICANT_FIELDS and conf < CONFIDENCE_INCLUDE_NO_FLAG:
            return f"[VERIFY: {val}]"
        return str(val)

    return re.sub(r"\{([^}]+)\}", replacer, template)


_ADL_ROW_DEFS = [
    ("Self Care Personal Hygiene", "Brushing Teeth", "p13_brushing_teeth_severity"),
    ("", "Flossing Teeth", "p13_flossing_teeth_severity"),
    ("Communication", "Speak for Extended periods of time", "p13_speak_extended_period_severity"),
    ("", "Speaking", "p13_speaking_difficulty_severity"),
    ("", "Indistinct Articulation", "p13_asked_to_repeat_themselves_severity"),
    ("Motor Function", "Mastication", "p13_mastication_severity"),
    ("", "Tasting", "p13_tasting_severity"),
    ("", "Swallowing", "p13_swallowing_severity"),
    ("", "Bruxism", "p13_bruxism_severity"),
    ("Sexual Function", "Kissing, Oral Activities", "p13_kissing_oral_activities_severity"),
]


def _build_adl_rows(
    section_rule: dict,
    all_fields: dict[str, Any],
    derived: dict[str, str],
    field_map: dict[str, str],
) -> list[list[str]]:
    """Build ADL table data rows from extraction fields."""
    rows = []
    for activity, example, fid in _ADL_ROW_DEFS:
        severity = _resolve_field(fid, all_fields, derived, field_map)
        difficulty = str(severity).strip() if severity else "None"
        rows.append([activity, example, difficulty])
    return rows


def _build_list_content(
    section_rule: dict,
    all_fields: dict[str, Any],
    derived: dict[str, str],
    field_map: dict[str, str],
) -> str:
    """Build list content from extraction data."""
    fid = section_rule.get("list_field_id")
    if not fid:
        return ""
    val = _resolve_field(fid, all_fields, derived, field_map)
    if isinstance(val, list):
        return "\n".join(f"• {item}" for item in val)
    return f"• {val}" if val else ""


# =============================================================================
# Derived fields
# =============================================================================

def _build_diagnosis_list(all_fields: dict[str, Any], schema_path: Path | None = None) -> str:
    """Build diagnosis list from schema fields marked is_diagnosis_code."""
    path = schema_path or _DEFAULT_SCHEMA_PATH
    if not path.exists():
        return ""
    schema = json.loads(path.read_text(encoding="utf-8"))
    dx_items: list[str] = []
    for page in schema.get("pages", []):
        for section in page.get("sections", []):
            for field in section.get("fields", []):
                if not field.get("is_diagnosis_code"):
                    continue
                fid = field["field_id"]
                label = field.get("field_label", fid)
                val = _unwrap_field_value(all_fields.get(fid))
                if val and str(val).strip().lower() not in ("no", "false", "none", ""):
                    code = field.get("icd_code", "")
                    desc = field.get("icd_description", label)
                    dx_items.append(f"{code} — {desc}" if code else desc)
    if dx_items:
        return "\n".join(f"{i+1}. {dx}" for i, dx in enumerate(dx_items))
    return ""


def _derive_fields(
    all_fields: dict[str, Any],
    field_map: dict[str, str],
    patient_name: str | None = None,
    patient_context: dict[str, str] | None = None,
) -> dict[str, str]:
    """Build derived/alias fields that rule templates expect but extraction doesn't provide directly."""
    derived: dict[str, str] = {}
    pctx = patient_context or {}

    def _get(fid: str) -> Any:
        return _resolve_field(fid, all_fields, derived, field_map)

    if pctx.get("patient_first_name") and pctx.get("patient_last_name"):
        derived["patient_first_name"] = pctx["patient_first_name"]
        derived["patient_last_name"] = pctx["patient_last_name"]
        derived["patient_name"] = f"{pctx['patient_first_name']} {pctx['patient_last_name']}"
        full_name = derived["patient_name"]
    else:
        full_name = patient_name or _get("p1_patient_name") or ""
        if full_name:
            parts = str(full_name).strip().split()
            derived["patient_first_name"] = parts[0] if parts else ""
            derived["patient_last_name"] = parts[-1] if len(parts) > 1 else ""
            derived["patient_name"] = str(full_name).strip()

    simple_aliases = {
        "patient_dob": ["p1_birth_date", "p1_dob"],
        "patient_sex": ["p1_gender", "p1_sex"],
        "patient_phone": ["p1_cell_phone", "p1_home_phone"],
        "patient_address": ["p1_address"],
        "occupation": ["p5_job_title"],
        "exam_date": ["p1_date"],
    }
    for target, sources in simple_aliases.items():
        if pctx.get(target):
            derived[target] = pctx[target]
            continue
        for src in sources:
            val = _get(src)
            if val is not None and str(val).strip():
                derived[target] = str(val)
                break

    gender = _get("p1_gender") or ""
    gender_lower = str(gender).strip().lower()
    if gender_lower in ("male", "m"):
        derived["patient_title"] = "Mr."
        derived["his_her"] = "his"
        derived["he_she"] = "he"
    elif gender_lower in ("female", "f"):
        derived["patient_title"] = "Ms."
        derived["his_her"] = "her"
        derived["he_she"] = "she"
    else:
        derived["patient_title"] = "Mr./Ms."
        derived["his_her"] = "his/her"
        derived["he_she"] = "he/she"

    vas_raw = _get("p14_max_opening_vas")
    if vas_raw is not None:
        try:
            vas_val = float(str(vas_raw).split("-")[0])
            if vas_val == 0:
                derived["pain_qualifier"] = "painless"
            elif vas_val <= 3:
                derived["pain_qualifier"] = "mildly painful"
            elif vas_val <= 6:
                derived["pain_qualifier"] = "moderately painful"
            else:
                derived["pain_qualifier"] = "severely painful"
        except (ValueError, TypeError):
            derived["pain_qualifier"] = "painful"
    else:
        derived["pain_qualifier"] = "painless"

    midline = _get("p15_midline_deviation")
    if midline and str(midline).strip().lower() not in ("none", "no", "centered", "0", ""):
        derived["midline_text"] = f"deviated to the {midline}"
    else:
        derived["midline_text"] = "centered"

    schema_dx = _build_diagnosis_list(all_fields)
    if schema_dx:
        derived["diagnosis_list"] = schema_dx
    else:
        dx_codes = {
            "p17_f45_8": ("F45.8", "Bruxism"),
            "p17_m79_1": ("M79.1", "Myalgia of Facial Muscles"),
            "p17_m65_80": ("M65.80", "Capsulitis/Inflammation of the TMJ"),
            "p17_g50_0": ("G50.0", "Trigeminal Nerve Neuropathic Pain"),
            "p17_m26_69_osteoarthritis": ("M26.69", "Osteoarthritis of the TMJ"),
            "p17_g51_0_halitosis": ("G51.0", "Halitosis"),
            "p17_k05_6": ("K05.6", "Gingival/Periodontal Disease"),
        }
        dx_items = []
        for fid, (code, desc) in dx_codes.items():
            val = _get(fid)
            if val and str(val).strip().lower() not in ("no", "false", "none", ""):
                dx_items.append(f"{code} — {desc}")
        if dx_items:
            derived["diagnosis_list"] = "\n".join(f"{i+1}. {dx}" for i, dx in enumerate(dx_items))

    epworth_fields = [
        ("Sitting and reading", "p11_epworth_sitting_reading", "epworth_sitting_reading"),
        ("Watching TV", "p11_epworth_watching_tv", "epworth_watching_tv"),
        ("Sitting inactive in a public place", "p11_epworth_sitting_inactive", "epworth_sitting_inactive"),
        ("Passenger in a car for 1 hour", "p11_epworth_passenger_car", "epworth_passenger_car"),
        ("Lying down in the afternoon", "p11_epworth_lying_down_afternoon", "epworth_lying_down_afternoon"),
        ("Sitting and talking to someone", "p11_epworth_sitting_talking", "epworth_sitting_talking"),
        ("Sitting quietly after lunch", "p11_epworth_sitting_after_lunch", "epworth_sitting_after_lunch"),
        ("In a car, stopped in traffic", "p11_epworth_car_traffic", "epworth_car_traffic"),
    ]
    epworth_items = []
    epworth_total = 0
    for label, canonical_fid, legacy_fid in epworth_fields:
        val = _get(canonical_fid) or _get(legacy_fid)
        if val is not None:
            try:
                score = int(str(val).strip())
                epworth_items.append(f"{label}: {score}")
                epworth_total += score
            except ValueError:
                pass
    if epworth_items:
        derived["epworth_activities"] = "\n".join(epworth_items)
        derived["epworth_sleepiness_scale_total_score"] = str(epworth_total)
    else:
        raw_total = _get("p11_epworth_total_score") or _get("Epworth_Sleepiness_Scale_Total_Score")
        if raw_total is not None:
            derived["epworth_sleepiness_scale_total_score"] = str(raw_total).strip()

    case_sources = {
        "case_number": ["case_number"],
        "claim_number": ["claim_number"],
        "venue": ["wcab_venue", "venue"],
        "date_of_injury": ["injury_date", "p5_date_of_injury", "date_of_injury"],
        "interpreter": ["interpreter_language", "interpreter"],
        "patient_city": ["patient_city"],
        "patient_state": ["patient_state"],
        "patient_zip": ["patient_zip"],
        "employer_name": ["employer_name"],
        "claims_admin_name": ["claims_admin_name"],
        "claims_admin_address": ["claims_admin_address"],
        "claims_admin_city": ["claims_admin_city"],
        "claims_admin_state": ["claims_admin_state"],
        "claims_admin_zip": ["claims_admin_zip"],
        "claims_admin_phone": ["claims_admin_phone"],
    }
    for target, sources in case_sources.items():
        if pctx.get(target):
            derived[target] = pctx[target]
            continue
        for src in sources:
            val = _get(src)
            if val is not None and str(val).strip():
                derived[target] = str(val)
                break

    for date_key in ("exam_date", "patient_dob", "date_of_injury"):
        raw = derived.get(date_key, "")
        m = re.match(r"(\d{1,2})/(\d{1,2})/(\d{4})", raw)
        if m:
            from datetime import datetime as _dt
            try:
                dt = _dt(int(m.group(3)), int(m.group(1)), int(m.group(2)))
                derived[f"{date_key}_formatted"] = dt.strftime("%B %d, %Y")
            except ValueError:
                pass

    addr = derived.get("patient_address", "")
    if addr:
        parts = [p.strip() for p in str(addr).split(",")]
        if len(parts) >= 3:
            derived.setdefault("patient_city", parts[-2].strip())
            last = parts[-1].strip().split()
            if len(last) >= 2:
                derived.setdefault("patient_state", last[0])
                derived.setdefault("patient_zip", last[1])
        elif len(parts) == 2:
            last = parts[-1].strip().split()
            if len(last) >= 2:
                derived.setdefault("patient_state", last[0])
                derived.setdefault("patient_zip", last[1])
        else:
            addr_parts = addr.split()
            if len(addr_parts) >= 3:
                derived.setdefault("patient_zip", addr_parts[-1] if addr_parts[-1].isdigit() else "")
                derived.setdefault("patient_state", addr_parts[-2] if len(addr_parts[-2]) == 2 else "")

    for muscle, vas_field in [
        ("p14_masseter_right_y_n", "p14_masseter_right_vas"),
        ("p14_masseter_left_y_n", "p14_masseter_left_vas"),
        ("p14_temporalis_right_y_n", "p14_temporalis_right_vas"),
        ("p14_temporalis_left_y_n", "p14_temporalis_left_vas"),
    ]:
        vas = _get(vas_field)
        if vas is not None:
            try:
                derived[muscle] = "Yes" if float(str(vas).split("-")[0]) > 0 else "No"
            except (ValueError, TypeError):
                derived[muscle] = "Yes" if str(vas).strip() else "No"

    derived.setdefault("additional_findings", "")
    return derived


# =============================================================================
# Narrative generation + hallucination guard
# =============================================================================

_PREAMBLE_PATTERNS = re.compile(
    r"^(here is|here's|based on|i will|i'll|the following|below is|let me|"
    r"sure,|certainly|of course|this section|the .* section)",
    re.IGNORECASE,
)

_REASONING_PATTERNS = re.compile(
    r"(^becomes$|^is not needed|^the correct response is|^the final answer is|"
    r"^the text is verified|^since ['\"]?p\d+_|^the generated text|^no changes|"
    r"^verified text:|^corrected text:|^output:$|^SECTION:\s|^GENERATED TEXT:)",
    re.IGNORECASE,
)

_INLINE_COMMENTARY = re.compile(
    r"\s*(is inconsistent with.*?[;.]|is an approximation.*?[;.]|should be understood that.*?[;.]|"
    r"therefore,?\s+it\s+should|note:.*?[;.]|this\s+is\s+because.*?[;.])",
    re.IGNORECASE,
)

_FIELD_ID_LINE = re.compile(r"['\"]?p\d+_\w+['\"]?\s*[:=]")


def _clean_narrative(text: str) -> str:
    """Strip LLM artifacts and verification reasoning from generated narrative text."""
    lines = text.strip().split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned.append(line)
            continue
        if _PREAMBLE_PATTERNS.match(stripped):
            continue
        if _REASONING_PATTERNS.match(stripped):
            continue
        if _FIELD_ID_LINE.search(stripped):
            continue
        if stripped.startswith("#"):
            line = stripped.lstrip("#").strip()
        line = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", line)
        line = _INLINE_COMMENTARY.sub("", line)
        if line.strip():
            cleaned.append(line)
    result = "\n".join(cleaned).strip()
    if result.startswith('"') and result.endswith('"'):
        result = result[1:-1]
    return result


def _generate_hybrid(
    section_rule: dict,
    all_fields: dict[str, Any],
    derived: dict[str, str],
    field_map: dict[str, str],
    patient_name: str,
    llm_clients: list[tuple[OpenAI, str]],
    _log: dict | None = None,
) -> str:
    """Fill deterministic parts of a hybrid template, send only [BRIDGE] gaps to LLM."""
    template = section_rule.get("hybrid_template", "")
    source_ids = section_rule.get("source_field_ids", [])

    relevant: dict[str, str] = {}
    for fid in source_ids:
        val = _resolve_field(fid, all_fields, derived, field_map)
        if val is not None:
            relevant[fid] = str(val)

    threshold = section_rule.get("min_required_fields", 1)
    if _log is not None:
        _log["fields_resolved_count"] = len(relevant)
        _log["fields_total_count"] = len(source_ids)
        _log["min_fields_threshold"] = threshold
        _log["min_fields_passed"] = len(relevant) >= threshold

    if len(relevant) < threshold:
        return ""

    filled = _fill_template(template, all_fields, derived, field_map)

    bridges = re.findall(r'\[BRIDGE:\s*(.+?)\]', filled)
    if not bridges or not llm_clients:
        result = re.sub(r'\[BRIDGE:\s*.+?\]', '', filled).strip()
        if _log is not None:
            _log["narrative_generated"] = bool(result)
            _log["narrative_raw_length"] = len(result)
            _log["narrative_clean_length"] = len(result)
        return result

    patient_last = derived.get("patient_last_name", patient_name.split()[-1] if patient_name else "")
    title_prefix = derived.get("patient_title", "Mr./Ms.")
    temperature = section_rule.get("temperature", 0.1)

    for bridge_instruction in bridges:
        context_window = filled[max(0, filled.index(f"[BRIDGE: {bridge_instruction}]") - 200):
                                 filled.index(f"[BRIDGE: {bridge_instruction}]") + 200]
        prompt = (
            f"You are writing a short connecting phrase for a clinical report about {title_prefix} {patient_last}.\n"
            f"Surrounding context:\n{context_window}\n\n"
            f"Instruction: {bridge_instruction}\n\n"
            f"Write ONLY the connecting phrase (10-30 words). No explanations."
        )
        bridge_text = ""
        for client, model in llm_clients:
            try:
                resp = client.chat.completions.create(
                    model=model, max_tokens=100, temperature=temperature,
                    messages=[{"role": "user", "content": prompt}],
                )
                bridge_text = resp.choices[0].message.content.strip()
                bridge_text = _clean_narrative(bridge_text)
                break
            except Exception:
                continue
        filled = filled.replace(f"[BRIDGE: {bridge_instruction}]", bridge_text, 1)

    result = filled.strip()
    if _log is not None:
        _log["narrative_generated"] = True
        _log["narrative_raw_length"] = len(result)
        _log["narrative_clean_length"] = len(result)
    return result


def _generate_narrative(
    section_rule: dict,
    all_fields: dict[str, Any],
    derived: dict[str, str],
    field_map: dict[str, str],
    patient_name: str,
    llm_clients: list[tuple[OpenAI, str]],
    _log: dict | None = None,
) -> str:
    """Generate narrative text using a full-context pass with confidence filtering.

    Fields below CONFIDENCE_UNCERTAIN (0.40) are excluded.
    Fields between 0.40-0.59 are annotated with [UNCERTAIN].
    """
    prompt = section_rule.get("generation_prompt", "")
    source_ids = section_rule.get("source_field_ids", [])
    examples = section_rule.get("few_shot_examples", [])

    relevant: dict[str, str] = {}
    for fid in source_ids:
        val = _resolve_field(fid, all_fields, derived, field_map)
        if val is not None:
            relevant[fid] = str(val)

    threshold = section_rule.get("min_required_fields", 1)

    if _log is not None:
        _log["fields_resolved_count"] = len(relevant)
        _log["fields_total_count"] = len(source_ids)
        _log["min_fields_threshold"] = threshold
        _log["min_fields_passed"] = len(relevant) >= threshold

    if len(relevant) < threshold:
        return ""

    context: dict[str, str] = {}
    for fid, fv in all_fields.items():
        conf = _get_confidence(fv)
        if conf < CONFIDENCE_UNCERTAIN:
            continue
        unwrapped = _unwrap_field_value(fv)
        if unwrapped is not None and str(unwrapped).strip():
            val_str = str(unwrapped)
            if conf < CONFIDENCE_INCLUDE_NO_FLAG:
                val_str = f"{val_str} [UNCERTAIN]"
            context[fid] = val_str
    context.update(derived)

    patient_last = derived.get("patient_last_name", patient_name.split()[-1] if patient_name else "")
    title = derived.get("patient_title", "Mr./Ms.")
    pronoun_he = derived.get("he_she", "he/she")
    pronoun_his = derived.get("his_her", "his/her")

    system_content = (
        f"You are an expert clinical report writer generating a section of an orofacial pain evaluation report.\n\n"
        f"SECTION: {section_rule.get('title', '')}\n"
        f"PATIENT: {patient_name}\n\n"
        f"INSTRUCTIONS:\n"
        f"- {prompt}\n"
        f"- Write ONLY the section content. Do NOT include any preamble, introduction, or meta-commentary.\n"
        f"- Do NOT start with phrases like 'Here is', 'Based on', 'I will', 'The following'.\n"
        f"- Use ONLY the patient name '{patient_name}' — never use names from examples.\n"
        f"- Write in third person using {title} {patient_last}.\n"
        f"- Use pronouns '{pronoun_he}' and '{pronoun_his}'.\n"
        f"- If a field value is missing or empty, omit that detail rather than inventing data.\n"
        f"- Fields marked [UNCERTAIN] should be referenced as 'reportedly' and recommended for verification.\n"
        f"- Be concise, professional, and use clinical terminology appropriate for a medical-legal report.\n"
    )

    messages: list[dict] = [{"role": "system", "content": system_content}]

    for ex in examples[:3]:
        inp = ex.get("input_fields", {})
        out = ex.get("output_text", "")
        if inp and out:
            out = out.replace("{patient_name}", patient_name)
            out = out.replace("{patient_last_name}", patient_last)
            out = out.replace("{patient_title}", title)
            out = out.replace("{his_her}", pronoun_his)
            out = out.replace("{he_she}", pronoun_he)
            messages.append({"role": "user", "content": f"Generate section from: {json.dumps(inp)}"})
            messages.append({"role": "assistant", "content": out})

    messages.append({
        "role": "user",
        "content": f"Generate section from: {json.dumps(context)}",
    })

    if not llm_clients:
        if _log is not None:
            _log["narrative_generated"] = False
            _log["narrative_raw_length"] = 0
            _log["narrative_clean_length"] = 0
        return f"[Narrative placeholder — LLM not configured]\nFields: {json.dumps(relevant, indent=2)}"

    last_error = None
    for client, model in llm_clients:
        try:
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                max_tokens=section_rule.get("max_tokens", 500),
                temperature=section_rule.get("temperature", 0.3),
            )
            raw = response.choices[0].message.content.strip()
            cleaned = _clean_narrative(raw)
            if _log is not None:
                _log["narrative_generated"] = True
                _log["narrative_raw_length"] = len(raw)
                _log["narrative_clean_length"] = len(cleaned)
            return cleaned
        except Exception as e:
            last_error = e
            logger.warning("LLM call failed for %s with %s: %s — trying next provider",
                           section_rule.get("section_id"), model, e)

    logger.error("All LLM providers failed for %s: %s", section_rule.get("section_id"), last_error)
    if _log is not None:
        _log["narrative_generated"] = False
        _log["narrative_raw_length"] = 0
        _log["narrative_clean_length"] = 0
    return f"[Generation failed: {last_error}]\nFields: {json.dumps(relevant)}"


def _verify_narrative(
    narrative_text: str,
    source_fields: dict[str, str],
    section_title: str,
    llm_clients: list[tuple[OpenAI, str]],
    _log: dict | None = None,
) -> str:
    """Self-RAG hallucination guard: verify every factual claim against source data."""
    if not narrative_text or not llm_clients:
        if _log is not None:
            _log["verified"] = False
            _log["verified_length"] = len(narrative_text) if narrative_text else 0
        return narrative_text

    prompt = (
        f"You generated this clinical report section:\n\n"
        f"SECTION: {section_title}\n"
        f"GENERATED TEXT: {narrative_text}\n\n"
        f"PATIENT-SPECIFIC DATA (values extracted from the patient's form):\n"
        f"{json.dumps(source_fields, indent=2)}\n\n"
        f"Verify the generated text. The rules are:\n"
        f"- KEEP all standard clinical boilerplate, medical explanations, anatomical descriptions, and procedural methodology text — these are expected in every report.\n"
        f"- KEEP all academic citations, journal references, and legal references verbatim.\n"
        f"- KEEP all standard threshold values (e.g. '400 Newtons', 'normal limits') — these are clinical constants, not patient data.\n"
        f"- CHECK patient-specific values (measurements, names, dates, scores) against the source data above. If a patient-specific value CONTRADICTS the source data, CORRECT it.\n"
        f"- REMOVE any patient-specific claims that are completely fabricated (not derivable from the source data).\n"
        f"- Do NOT remove medical explanations, test descriptions, or citations just because they aren't in the source data — they are standard clinical content.\n\n"
        f"CRITICAL: Output ONLY the final verified clinical text, nothing else. "
        f"Do NOT include any reasoning, explanations, field names, comparisons, "
        f"intermediate steps, or commentary. Do NOT write words like 'becomes', "
        f"'the correct response is', 'is not needed', or 'Since'. "
        f"Your entire response must read as a polished clinical report paragraph."
    )

    verify_max_tokens = max(1500, len(narrative_text) // 3)

    for client, model in llm_clients:
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=verify_max_tokens,
                temperature=0.0,
            )
            verified = response.choices[0].message.content.strip()
            if verified:
                result = _clean_narrative(verified)
                if len(result) < len(narrative_text) * 0.3:
                    logger.warning(
                        "Verification stripped >70%% of '%s' (%d -> %d chars) — keeping unverified",
                        section_title, len(narrative_text), len(result),
                    )
                    if _log is not None:
                        _log["verified"] = True
                        _log["verified_length"] = len(narrative_text)
                        _log["verification_reverted"] = True
                    return narrative_text
                if _log is not None:
                    _log["verified"] = True
                    _log["verified_length"] = len(result)
                return result
        except Exception as e:
            logger.warning("Verification failed for '%s' with %s: %s", section_title, model, e)

    if _log is not None:
        _log["verified"] = False
        _log["verified_length"] = len(narrative_text)
    return narrative_text


# =============================================================================
# Cover page + preamble builder
# =============================================================================

_SUBSECTION_HEADERS = {
    "headaches", "facial pain", "temporomandibular joint", "dental",
    "sleep disturbances", "other relating complaints are:",
}


def _render_narrative_text(doc: DocxDocument, text: str) -> None:
    """Render narrative text with proper formatting: subsection headers, bullets, paragraphs."""
    for para in text.split("\n"):
        stripped = para.strip()
        if not stripped:
            continue

        is_bullet = stripped.startswith("\u2022") or stripped.startswith("- ") or stripped.startswith("* ")
        clean_text = stripped.lstrip("\u2022-* ").strip() if is_bullet else stripped

        is_subsection_header = stripped.lower().rstrip(":") in _SUBSECTION_HEADERS or stripped.lower() in _SUBSECTION_HEADERS

        if is_subsection_header:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif is_bullet:
            p = doc.add_paragraph(clean_text, style="List Bullet")
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
        else:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)


def _add_title_headings(doc: DocxDocument) -> None:
    """Add the two centered bold underlined title lines."""
    for text in ("INITIAL REPORT IN THE FIELD OF DENTISTRY",
                 "AND REQUEST FOR AUTHORIZATION"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.underline = True
        run.font.name = "Arial"
        run.font.size = Pt(11)


def _set_cell_border(cell, **kwargs):
    """Set borders on a table cell. kwargs: top, bottom, left, right with values like '0' or '4'."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = tc_pr.makeelement(qn("w:tcBorders"), {})
        tc_pr.append(borders)
    for edge, val in kwargs.items():
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = borders.makeelement(qn(f"w:{edge}"), {})
            borders.append(el)
        el.set(qn("w:val"), "single" if val != "0" else "none")
        el.set(qn("w:sz"), val)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def _add_field_line(cell, label: str, value: str) -> None:
    """Append a 'Label:\\tValue' line to a table cell."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{label}\t{value}")
    run.font.name = "Arial"
    run.font.size = Pt(10)


def _add_section_header(cell, text: str) -> None:
    """Add a bold section label (Patient, Claims Administrator, etc.) to a cell."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)


def _build_cover_page(
    doc: DocxDocument,
    derived: dict[str, str],
    patient_context: dict[str, str] | None = None,
) -> None:
    """Build cover sheet (page 1) and preamble (page 2) matching training reports."""
    pctx = patient_context or {}

    def _val(key: str) -> str:
        return derived.get(key, "") or pctx.get(key, "")

    # -- Page 1: Cover sheet --------------------------------------------------
    _add_title_headings(doc)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(4.2)
    table.columns[1].width = Inches(2.3)

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    _set_cell_border(left_cell, top="0", bottom="0", left="0", right="0")
    _set_cell_border(right_cell, top="4", bottom="4", left="4", right="4")

    # Left column: case info
    _add_field_line(left_cell, "DATE:", _val("exam_date"))
    _add_field_line(left_cell, "CLAIM:", _val("claim_number"))
    _add_field_line(left_cell, "WCAB:", _val("venue"))
    _add_field_line(left_cell, "Case #:", _val("case_number"))
    _add_field_line(left_cell, "Date of Injury:", _val("date_of_injury"))
    _add_field_line(left_cell, "Date of current exam:", _val("exam_date"))
    _add_field_line(left_cell, "Interpreter:", _val("interpreter"))

    # Left column: patient info
    _add_section_header(left_cell, "Patient")
    _add_field_line(left_cell, "Last Name:", _val("patient_last_name"))
    _add_field_line(left_cell, "First Name:", _val("patient_first_name"))
    _add_field_line(left_cell, "Sex:", _val("patient_sex"))
    _add_field_line(left_cell, "Date of Birth:", _val("patient_dob"))
    _add_field_line(left_cell, "Occupation:", _val("occupation"))
    _add_field_line(left_cell, "Address:", _val("patient_address"))
    _add_field_line(left_cell, "City:", _val("patient_city"))
    _add_field_line(left_cell, "State:", _val("patient_state"))
    _add_field_line(left_cell, "Zip:", _val("patient_zip"))
    _add_field_line(left_cell, "Phone No.:", _val("patient_phone"))

    # Left column: claims administrator
    _add_section_header(left_cell, "Claims Administrator/Insurer")
    _add_field_line(left_cell, "Name:", pctx.get("claims_admin_name", ""))
    _add_field_line(left_cell, "Address:", pctx.get("claims_admin_address", ""))
    _add_field_line(left_cell, "City:", pctx.get("claims_admin_city", ""))
    _add_field_line(left_cell, "State:", pctx.get("claims_admin_state", ""))
    _add_field_line(left_cell, "Zip:", pctx.get("claims_admin_zip", ""))
    _add_field_line(left_cell, "Phone:", pctx.get("claims_admin_phone", ""))

    # Left column: employer
    _add_section_header(left_cell, "Employer")
    _add_field_line(left_cell, "Name:", pctx.get("employer_name", ""))
    _add_field_line(left_cell, "Address:", pctx.get("employer_address", ""))
    _add_field_line(left_cell, "City:", pctx.get("employer_city", ""))
    _add_field_line(left_cell, "State:", pctx.get("employer_state", ""))
    _add_field_line(left_cell, "Zip:", pctx.get("employer_zip", ""))

    # Right column: legal demand box
    for para in right_cell.paragraphs:
        right_cell._tc.remove(para._element)
    p = right_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(_LEGAL_DEMAND_BOX)
    run.font.name = "Arial"
    run.font.size = Pt(7)

    doc.add_page_break()

    # -- Page 2: Preamble -----------------------------------------------------
    _add_title_headings(doc)

    # Disclaimer (all-caps, justified)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(_PREAMBLE_DISCLAIMER)
    run.font.name = "Arial"
    run.font.size = Pt(11)

    # Exam intro (with patient name insertion)
    intro_text = _PREAMBLE_EXAM_INTRO.format(
        patient_title=_val("patient_title") or "Mr./Ms.",
        patient_last_name=_val("patient_last_name") or "Patient",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(intro_text)
    run.font.name = "Arial"
    run.font.size = Pt(11)

    # Conclusions basis
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(_PREAMBLE_CONCLUSIONS_BASIS)
    run.font.name = "Arial"
    run.font.size = Pt(11)

    # Records demand (italic)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(_PREAMBLE_RECORDS_DEMAND)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(11)

    # Reserve rights
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(_PREAMBLE_RESERVE_RIGHTS)
    run.font.name = "Arial"
    run.font.size = Pt(11)

    # Conditional interpreter line
    interpreter_name = _val("interpreter")
    if interpreter_name.strip():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"{interpreter_name} served as an interpreter during this examination.")
        run.font.name = "Arial"
        run.font.size = Pt(11)

    doc.add_page_break()


# =============================================================================
# Main report generation
# =============================================================================

def generate_clinical_report(
    extraction_data: dict[str, Any],
    patient_context: dict[str, str] | None = None,
    rules_path: Path | None = None,
) -> tuple[bytes, list[dict], dict]:
    """Generate a clinical DOCX report from extraction data using learned rules.

    Args:
        extraction_data: Extraction result with pages[].field_values
        patient_context: Optional case demographics from patient_medical_history
        rules_path: Override path to report_rules.json

    Returns:
        Tuple of (DOCX file bytes, low_confidence_fields list, generation_log dict)
    """
    rules = _load_rules(rules_path)
    sections = rules.get("sections", [])
    global_fmt = rules.get("global_formatting", {})

    all_fields = _flatten_all_fields(extraction_data.get("pages", []))
    field_map = _FIELD_MAP

    patient_name = (
        extraction_data.get("patient_name")
        or _resolve_field("p1_patient_name", all_fields, {}, field_map)
        or "Patient"
    )
    if isinstance(patient_name, str):
        patient_name = patient_name.strip()

    derived = _derive_fields(all_fields, field_map, patient_name, patient_context)
    llm_clients = _build_llm_clients()
    low_confidence = _collect_low_confidence_fields(all_fields)

    if llm_clients:
        logger.info("LLM providers configured: %s", [m for _, m in llm_clients])
    else:
        logger.warning("No LLM clients available — narrative sections will use placeholders")

    doc = DocxDocument()

    style = doc.styles["Normal"]
    font = style.font
    font.name = global_fmt.get("font_name", "Arial")
    font.size = Pt(global_fmt.get("font_size", 11))

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    _build_cover_page(doc, derived, patient_context)

    generation_log: dict[str, dict] = {}

    for sec_rule in sections:
        section_id = sec_rule.get("section_id", "")
        content_type = sec_rule.get("content_type", "")
        conditions = sec_rule.get("conditions", [])
        slog: dict[str, Any] = {
            "content_type": content_type,
            "source_field_ids": sec_rule.get("source_field_ids", []),
            "has_generation_prompt": bool(sec_rule.get("generation_prompt")),
            "has_few_shot": bool(sec_rule.get("few_shot_examples")),
        }

        if section_id in _COVER_SKIP_SECTION_IDS:
            slog["skipped_cover"] = True
            slog["conditions_result"] = True
            slog["final_text_length"] = 0
            generation_log[section_id] = slog
            continue
        slog["skipped_cover"] = False

        cond_result = not conditions or _check_conditions(conditions, all_fields, derived, field_map)
        slog["conditions_result"] = cond_result
        if not cond_result:
            slog["final_text_length"] = 0
            generation_log[section_id] = slog
            continue

        title = sec_rule.get("title", "")
        if title and content_type != "static_text":
            heading = doc.add_heading(title, level=2)
            for run in heading.runs:
                run.font.name = "Arial"
                run.font.size = Pt(14)
                run.bold = True
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        final_text = ""

        if content_type == "static_text":
            text = sec_rule.get("static_content", "")
            if text:
                if title:
                    heading = doc.add_heading(title, level=2)
                    for run in heading.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(14)
                        run.bold = True
                        run.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                for line in text.split("\n"):
                    p = doc.add_paragraph(line)
                    p.style.font.name = "Arial"
                    p.style.font.size = Pt(11)
            final_text = text or ""

        elif content_type in ("direct_fill", "formatted_fill"):
            template = sec_rule.get("template", "")
            if template:
                filled = _fill_template(template, all_fields, derived, field_map)
                for line in filled.split("\n"):
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(2)
                final_text = filled
            slog["template_filled"] = bool(template)

        elif content_type == "narrative":
            if sec_rule.get("hybrid_template"):
                text = _generate_hybrid(sec_rule, all_fields, derived, field_map, patient_name, llm_clients, _log=slog)
            else:
                text = _generate_narrative(sec_rule, all_fields, derived, field_map, patient_name, llm_clients, _log=slog)
            if text:
                # Strip duplicate heading if the generated text starts with the section title
                if title:
                    for prefix in (title, title.upper(), title.title()):
                        if text.strip().startswith(prefix):
                            text = text.strip()[len(prefix):].lstrip("\n").lstrip()
                            break

                # Skip verification for pure hybrid templates or sections that opt out
                hybrid_tmpl = sec_rule.get("hybrid_template") or ""
                is_pure_hybrid = hybrid_tmpl and "[BRIDGE:" not in hybrid_tmpl
                skip_verify = sec_rule.get("skip_verification", False)
                if not is_pure_hybrid and not skip_verify:
                    source_ids = sec_rule.get("source_field_ids", [])
                    relevant = {fid: str(v) for fid in source_ids
                                if (v := _resolve_field(fid, all_fields, derived, field_map)) is not None}
                    text = _verify_narrative(text, relevant, sec_rule.get("title", ""), llm_clients, _log=slog)

                _render_narrative_text(doc, text)
            final_text = text or ""

        elif content_type == "list":
            text = _build_list_content(sec_rule, all_fields, derived, field_map)
            if text:
                for line in text.split("\n"):
                    doc.add_paragraph(line, style="List Bullet")
            final_text = text or ""

        elif content_type == "table":
            columns = sec_rule.get("table_columns", [])
            if columns:
                table_rows = _build_adl_rows(sec_rule, all_fields, derived, field_map)
                table = doc.add_table(rows=1 + len(table_rows), cols=len(columns))
                table.style = "Table Grid"
                for i, col in enumerate(columns):
                    cell = table.rows[0].cells[i]
                    cell.text = col.get("header", "")
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                for row_idx, row_data in enumerate(table_rows, start=1):
                    for col_idx, val in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = val
                final_text = str(table_rows)

        elif content_type == "conditional_block":
            child_sections = sec_rule.get("child_sections", [])
            for child in child_sections:
                child_conds = child.get("conditions", [])
                if child_conds and not _check_conditions(child_conds, all_fields, derived, field_map):
                    continue
                child_text = child.get("static_content") or ""
                if child_text:
                    doc.add_paragraph(child_text)
                    final_text += child_text + "\n"

        slog["final_text_length"] = len(final_text)
        generation_log[section_id] = slog

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), low_confidence, generation_log
